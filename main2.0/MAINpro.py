from flask import Flask, request, render_template,send_file
import re
import PyPDF2
import os
from docx import Document
import win32com.client as win32
import os
from bs4 import BeautifulSoup
import sqlite3
import pyodbc


txt_path = 'txt.txt'

def get_all_files(path):
    file_names = []
    for root, dirs, files in os.walk(path):
        for file in files:
            file_names.append(file)
    return file_names


#............pdf转txt
#......................................
def pdf_txt(document_path,txt_path):
#打开PDF文件
    pdf_file = open(document_path, 'rb')

# 创建PDF文件阅读器对象
    pdf_reader = PyPDF2.PdfReader(pdf_file)

# 获取PDF文档的页数
    pages_num = len(pdf_reader.pages)

# 将PDF文档转换为TXT格式
    with open(txt_path, 'w') as txt_file:
        for page_index in range(pages_num):
            page_content = pdf_reader.pages[page_index].extract_text()
            txt_file.write(page_content)

# 关闭PDF文件
    pdf_file.close()
    return txt_path

# #.........word转txt

def word_txt(document_path, txt_path):
    # 使用python-docx库加载Word文档
    doc = Document(document_path)

    # 打开txt文件，准备写入内容
    with open(txt_path, 'w') as txt_file:
        # 遍历文档的每个段落
        for para in doc.paragraphs:
            line = para.text.strip()  # 获取段落的文本内容，并去除首尾空格
            if line:  # 如果内容不为空，则写入txt文件
                txt_file.write(line + '\n')

    return txt_path


# def word_txt(document_path,txt_path):
# # 打开DOCX文件
#     doc_path = document_path
#     doc = Document(doc_path)

# # 保存为临时的DOC文件
#     temp_doc_path = 'temp.doc'
#     doc.save(temp_doc_path)

# # 指定输出PDF文件路径
#     pdf_path = document_path.replace('.docx','.pdf')

# # 创建Word应用程序实例
#     word = win32.gencache.EnsureDispatch('Word.Application')

#     try:
#     # 打开临时的DOC文件
#         doc_word = word.Documents.Open(os.path.abspath(temp_doc_path))

#     # 将DOC文件另存为PDF
#         doc_word.SaveAs(os.path.abspath(pdf_path), FileFormat=17)

#     finally:
#     # 关闭Word应用程序并退出
#         if doc_word:
#             doc_word.Close()
#         word.Quit()

# # 删除临时的DOC文件
#     os.remove(temp_doc_path)

#     print(f"DOCX文件已成功转换为PDF：{pdf_path}")
#     pdf_txt(pdf_path,txt_path)
#     return txt_path



# def txt_search(path,key_word):
#     for f_name in get_all_files(path):
#         document_path = path+"\\"+f_name
#         index = document_path.find('.pdf')
#         if index != -1:
#             f = open(txt_path,'r')
#             i = 1
#             l1 = []
#             ls = str(document_path)
#             l1.append(ls) 
#             lines = f.readlines()
#             for line in lines:
#                 if re.search(key_word,line):
#                     line1 = f'Line {i,line}'
#                     l1.append(line1)
#                 i = i+1
#             # txt_search(pdf_txt(document_path,txt_path),name)
#         index2 = path.find('.docx')
#         if index2 != -1:
#             f = open(txt_path,'r')
#             i = 1
#             l1 = []
#             ls = str(document_path)
#             l1.append(ls) 
#             lines = f.readlines()
#             for line in lines:
#                 if re.search(key_word,line):
#                     line1 = f'Line {i,line}'
#                     l1.append(line1)
#                 i = i+1

#     return l1

app = Flask(__name__)

testpath = ''
l1 = []
str_s = '-----------------------------'
@app.route('/')
def index():
    global testpath
    directory_path = 'C:/'  # 本地文件目录的路径

    # 获取当前路径和上级路径
    current_path = request.args.get('path', '')
    parent_path = os.path.dirname(current_path)
    testpath = current_path
    # 拼接完整路径
    if current_path:
        directory_path = os.path.join(directory_path, current_path)

    # 遍历目录下的文件和文件夹
    file_list = []
    for item in os.listdir(directory_path):
        item_path = os.path.join(directory_path, item)

        if os.path.isfile(item_path):
            file_list.append({
                'name': item,
                'type': 'file'
            })
        elif os.path.isdir(item_path):
            file_list.append({
                'name': item,
                'type': 'directory',
                'path': os.path.join(current_path, item)  # 记录子目录的路径
            })

    return render_template('test3.html', current_path=current_path, parent_path=parent_path, file_list=file_list)
#建立数据库
# db_file = 'index.db'

# def index_folder():
#     folder_path = testpath

#     # 连接到数据库或创建新的数据库文件
#     conn = sqlite3.connect(db_file)

#     # 创建文件索引表
#     conn.execute('''CREATE TABLE IF NOT EXISTS file_index
#                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
#                     filename TEXT,
#                     path TEXT);''')

#     # 遍历文件夹并索引匹配的文件
#     for root, dirs, files in os.walk(folder_path):
#         for filename in files:
#             if filename.endswith('.pdf') or filename.endswith('.docx'):
#                 file_path = os.path.join(root, filename)

#                 # 将文件信息插入数据库
#                 conn.execute('INSERT INTO file_index (filename, path) VALUES (?, ?)', (filename, file_path))

#     # 提交并关闭数据库连接
#     conn.commit()
#     conn.close()

#     return '索引建立成功！'

server = 'localhost'
database = 'shortterm'
username = 'sa'
password = 'wxLk1009.'
print(testpath)
# 建立数据库连接
conn = pyodbc.connect(f'DRIVER={{SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}')

def index_folder(testpath):
    folder_path = testpath

    # 创建文件索引表
    cursor = conn.cursor()
    cursor.execute('''IF NOT EXISTS (SELECT * FROM sys.tables WHERE name='file_index')
                      CREATE TABLE file_index (
                          id INT IDENTITY(1,1) PRIMARY KEY,
                          filename varchar(100),
                          path varchar(100)
                      );''')

    # 遍历文件夹并索引匹配的文件
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            print(1111)
            if filename.endswith('.pdf') or filename.endswith('.docx'):
                file_path = os.path.join(root, filename)
                print(file_path)
                print(filename)
                # 将文件信息插入数据库
                cursor.execute('INSERT INTO file_index (filename, path) VALUES (?, ?)', (filename, file_path))

    # 提交并关闭数据库连接
    cursor.commit()
    cursor.close()

    return '索引建立成功！'

# 调用索引函数
index_folder(testpath)

# 关闭数据库连接
conn.close()

@app.route('/center/add')
def center():
        path = "C:\\"+testpath
        name = request.args.get('key_word')  # args取get方式参数
        

        for f_name in get_all_files(path):
            document_path = path+"\\"+f_name


            #判断是否为pdf文件
            index = document_path.find('.pdf')
            if index != -1:
                f = open(txt_path,'r')
                key_word = name
                i = 1
                bol = 0
                lines = f.readlines()
                for line in lines:
                    if re.search(key_word,line):
                        bol = 1
                        break
                if bol == 1:
                    l1.append(document_path)
                    l1.append(str_s)
                pdf_txt(document_path,txt_path)
                for line in lines:
                    if re.search(key_word,line):
                        ls = f'line{i,line}'
                        l1.append(ls)
                    i = i+1
            
            #判断是否为word文件

            index2 = document_path.find('.docx')
            if index2 != -1:
                f = open(txt_path,'r')
                key_word = name
                i = 1
                bol = 0
                lines = f.readlines()
                for line in lines:
                    if re.search(key_word,line):
                        bol = 1
                        break
                if bol == 1:
                    l1.append(document_path)
                    l1.append(str_s)
                word_txt(document_path,txt_path)
                for line in lines:
                    if re.search(key_word,line):
                        ls = f'line{i,line}'
                        l1.append(ls)
                    i = i+1
        # return "路径：%s\n ************\n 行号：%d 对应关键行：%s" % (document_path, i, line)
        return render_template('test3.html',lines = l1,data = name)

@app.route('/download')
def download():
    # 获取需要保存的数据（document_path、i和line）并进行处理
    # 将需要保存的数据拼接成一个字符串
    file_name = "search_results.txt"  # 下载的文件名

    # 将数据保存到文件中
    with open(file_name, "w") as file:
        for line in l1:
            file.write(line)
            file.write('\n')

    # 返回文件给用户下载
    return send_file(file_name, as_attachment=True)


if __name__ == '__main__':
    app.run()
